import requests as rq
from bs4 import BeautifulSoup
import docx
from docx.shared import Cm
# import time

# def sleeptime(hour,min,sec):
#     return hour*3600 + min*60 + sec

imgfolder = "C:\\Users\\2200432\\Desktop\\python\\Baha\\image"
baha_url = 'https://gnn.gamer.com.tw/'
cellphone_game = "index.php?k=4"
pc_game = "index.php?k=1"
console_game = "index.php?k=3"
anime = "index.php?k=5"
none = ""


def catch(select, name):
	i = 1
	r = rq.get(baha_url + select)
	print("\nPlease wait.......")
	if r.status_code == rq.codes.ok:
		doc.add_heading('Gnn_News', level=1)
		soup = BeautifulSoup(r.text, "html.parser")
		url = soup.find_all("div", class_="GN-lbox2B", limit=25)
		for p in url:
			img_r = rq.get(p.select_one("img").get("src")).content
			with open('%s\\%s.jpg' % (imgfolder, i), 'wb') as f:
				f.write(img_r)
			title_text = p.find("h1")
			words = p.find("p")

			page_url = "https:" + p.select_one("a").get("href")
			img_url = p.select_one("img").get("src")
			title = title_text.select_one("a").getText()
			text = words.getText()[:-8]
			doc.add_heading("第" + str(i) + "則", level=2)
			doc.add_heading(title, level=2)
			doc.add_picture(imgfolder + "\\" + str(i) + ".jpg", width=Cm(10))
			doc.add_heading("URL", level=2)
			doc.add_paragraph(page_url)
			doc.add_heading('內文', level=2)
			doc.add_paragraph(text)
			doc.add_page_break()
			# data.append([page_url, img_url, title, text])
			i += 1
		doc.save(name + '_news.docx')
		print("Finish")
	else:
		print("網址or網頁資料有問題")



if __name__	== '__main__':
	# data = []
	str_name = ["Cellphone_Game", "PC_Game", "Console_Game", "Anime", "Total_Gnn"]
	doc = docx.Document()
	print("What do you want?")
	user_select = int(input("1 : cellphone_game, 2 : pc_game, 3 : console_game, 4 : anime, 5 : total_gnn ==> "))
	if(user_select == 1):
		catch(cellphone_game, str_name[user_select - 1])
	elif(user_select == 2):
		catch(pc_game, str_name[user_select - 1])
	elif(user_select == 3):
		catch(console_game, str_name[user_select - 1])
	elif(user_select == 4):
		catch(anime, str_name[user_select - 1])
	elif(user_select == 5):
		catch(none, str_name[user_select - 1])
	else:
		print("input error")


